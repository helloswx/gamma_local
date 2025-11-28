"""
Gamma API 调用工具
将文本内容或 PDF 发送给 Gamma，生成 PPT 并导出为 PDF
"""
import os
import time
import json
import hashlib
import requests
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
import PyPDF2
from datetime import datetime

# Selenium 相关导入（可选，用于浏览器自动化）
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.chrome.options import Options
    from webdriver_manager.chrome import ChromeDriverManager
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False
    print("警告: Selenium 未安装，浏览器自动化功能将不可用")
    print("如需使用浏览器自动化导出，请运行: pip install selenium webdriver-manager")

# 加载环境变量
load_dotenv()

# Gamma API 配置
GAMMA_API_KEY = os.getenv("GAMMA_API_KEY")
GAMMA_THEME_ID = os.getenv("GAMMA_THEME_ID")  # 可选：如果未设置则使用工作区默认主题
GAMMA_API_BASE_URL = "https://public-api.gamma.app/v1.0"

# 目录配置
DATASET_DIR = Path("dataset")
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

# 记录文件路径
RECORDS_FILE = Path("generation_records.json")


def extract_text_from_docx(file_path):
    """
    从 .docx 文件中提取文本内容和图片链接
    
    返回: (text_content, image_urls)
    """
    try:
        doc = Document(file_path)
        text_parts = []
        image_urls = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        # 从所有文本中提取图片 URL（包括段落和表格）
        import re
        all_text = "\n".join(text_parts)
        # 查找所有 http/https 链接
        urls = re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', all_text)
        for url in urls:
            # 检查是否是图片 URL（通过扩展名或常见图片服务域名）
            url_lower = url.lower()
            is_image = (
                any(url_lower.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.bmp']) or
                any(domain in url_lower for domain in ['imgur.com', 'imgbb.com', 'cloudinary.com', 'unsplash.com', 'pexels.com'])
            )
            if is_image and url not in image_urls:
                image_urls.append(url)
        
        text_content = "\n".join(text_parts)
        return text_content, image_urls
    except Exception as e:
        print(f"读取 .docx 文件时出错: {e}")
        return None, []


def load_records():
    """加载生成记录"""
    if RECORDS_FILE.exists():
        try:
            with open(RECORDS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"加载记录文件失败: {e}")
            return {}
    return {}


def save_records(records):
    """保存生成记录"""
    try:
        with open(RECORDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"保存记录文件失败: {e}")
        return False


def get_file_hash(file_path):
    """计算文件的哈希值，用于唯一标识文件"""
    try:
        hash_md5 = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except Exception as e:
        print(f"计算文件哈希失败: {e}")
        return None


def check_existing_generation(file_path):
    """
    检查文件是否已经生成过
    
    返回: (exists, record) 如果存在则返回记录，否则返回 None
    """
    file_hash = get_file_hash(file_path)
    if not file_hash:
        return False, None
    
    records = load_records()
    
    # 检查文件路径和哈希值
    file_path_str = str(Path(file_path).resolve())
    
    for record_id, record in records.items():
        if (record.get('file_path') == file_path_str or 
            record.get('file_hash') == file_hash):
            return True, record
    
    return False, None


def add_generation_record(file_path, generation_id, gamma_url, status="completed"):
    """
    添加生成记录
    
    参数:
    - file_path: 输入文件路径
    - generation_id: 生成任务 ID
    - gamma_url: Gamma 演示文稿 URL
    - status: 生成状态
    """
    file_hash = get_file_hash(file_path)
    if not file_hash:
        return False
    
    records = load_records()
    
    # 使用 generation_id 作为记录 ID
    record = {
        "file_path": str(Path(file_path).resolve()),
        "file_name": Path(file_path).name,
        "file_hash": file_hash,
        "generation_id": generation_id,
        "gamma_url": gamma_url,
        "status": status,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    
    records[generation_id] = record
    return save_records(records)


def update_generation_record(generation_id, **kwargs):
    """更新生成记录"""
    records = load_records()
    
    if generation_id in records:
        records[generation_id].update(kwargs)
        records[generation_id]["updated_at"] = datetime.now().isoformat()
        return save_records(records)
    
    return False


def list_generations():
    """列出所有生成记录"""
    records = load_records()
    
    if not records:
        print("没有找到生成记录")
        return
    
    print("\n" + "=" * 80)
    print("已生成的文档记录")
    print("=" * 80)
    
    for i, (record_id, record) in enumerate(sorted(records.items(), 
                                                    key=lambda x: x[1].get('created_at', ''), 
                                                    reverse=True), 1):
        print(f"\n记录 #{i}")
        print(f"  文件名: {record.get('file_name', 'N/A')}")
        print(f"  生成 ID: {record.get('generation_id', 'N/A')}")
        print(f"  状态: {record.get('status', 'N/A')}")
        print(f"  Gamma URL: {record.get('gamma_url', 'N/A')}")
        print(f"  创建时间: {record.get('created_at', 'N/A')}")
        print(f"  更新时间: {record.get('updated_at', 'N/A')}")
    
    print("\n" + "=" * 80)


def extract_text_from_pdf(file_path):
    """从 PDF 文件中提取文本内容"""
    try:
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        print(f"读取 PDF 文件时出错: {e}")
        return None


def generate_presentation(input_text, theme_id=None, num_cards=None, additional_instructions=None, has_image_urls=False):
    """
    调用 Gamma API 生成演示文稿
    
    参数说明：
    - input_text: 输入文本内容
    - theme_id: 主题 ID（可选），如果不提供则使用工作区默认主题
    - num_cards: 幻灯片数量（可选）
    - additional_instructions: 额外指示（可选）
    """
    if not GAMMA_API_KEY:
        raise ValueError("请设置 GAMMA_API_KEY 环境变量")
    
    url = f"{GAMMA_API_BASE_URL}/generations"
    headers = {
        "X-API-KEY": GAMMA_API_KEY,
        "Content-Type": "application/json"
    }
    
    # 构建请求参数（必需参数）
    payload = {
        "inputText": input_text,
        "textMode": "generate",
        "format": "presentation",
        "cardSplit": "auto",
        "exportAs": "pdf"
    }
    
    # 可选参数：themeId（根据官方文档，如果不提供则使用工作区默认主题）
    if theme_id:
        payload["themeId"] = theme_id
    
    # 可选参数
    if num_cards:
        payload["numCards"] = num_cards
    
    if additional_instructions:
        payload["additionalInstructions"] = additional_instructions
    
    # 文本选项（根据官方文档：https://developers.gamma.app/docs/generate-api-parameters-explained）
    payload["textOptions"] = {
        "amount": "detailed",  # 可选: "brief", "medium", "detailed"
        "tone": "professional",  # 可选: 一个或多个词，1-500字符
        "audience": "general",  # 可选: 一个或多个词，1-500字符，描述目标受众
        "language": "en"  # 可选: 默认 "en"，支持多种语言
    }
    
    # 图像选项（根据官方文档）
    # 如果提供了图片 URL，使用 noImages 以仅使用提供的图片（根据官方文档）
    # 否则使用 aiGenerated 让 Gamma 生成图片
    image_source = "noImages" if has_image_urls else "aiGenerated"
    
    payload["imageOptions"] = {
        "source": image_source,  # 可选: "aiGenerated"(默认), "pictographic", "unsplash", "giphy", "webAllImages", "webFreeToUse", "webFreeToUseCommercially", "placeholder", "noImages"
    }
    
    # 仅在 source 为 "aiGenerated" 时添加 model 和 style
    if image_source == "aiGenerated":
        payload["imageOptions"]["model"] = "imagen-4-pro"  # 可选: 仅在 source 为 "aiGenerated" 时相关
        payload["imageOptions"]["style"] = "photorealistic"  # 可选: 仅在 source 为 "aiGenerated" 时相关
    
    # 卡片选项（根据官方文档）
    # cardOptions 支持 dimensions 和 headerFooter，不支持 size、includeHeader、includeFooter
    payload["cardOptions"] = {
        "dimensions": "fluid"  # 可选: "fluid"(默认), "16x9", "4x3" (format 为 presentation 时)
        # headerFooter 是可选的，如果需要可以添加，例如：
        # "headerFooter": {
        #     "topRight": {
        #         "type": "image",
        #         "source": "themeLogo",
        #         "size": "sm"
        #     },
        #     "bottomRight": {
        #         "type": "cardNumber"
        #     },
        #     "hideFromFirstCard": true,
        #     "hideFromLastCard": false
        # }
    }
    
    print("正在调用 Gamma API 生成演示文稿...")
    response = requests.post(url, json=payload, headers=headers)
    
    # 201 (Created) 和 200 (OK) 都表示成功
    if response.status_code not in [200, 201]:
        print(f"API 调用失败: {response.status_code}")
        print(f"错误信息: {response.text}")
        return None
    
    result = response.json()
    generation_id = result.get("generationId")
    
    if not generation_id:
        print(f"未获取到 generationId: {result}")
        return None
    
    print(f"生成任务已创建，generationId: {generation_id}")
    return generation_id


def check_generation_status(generation_id):
    """
    检查生成状态
    
    返回: (status, result_data)
    status: "pending", "processing", "completed", "failed"
    """
    if not GAMMA_API_KEY:
        raise ValueError("请设置 GAMMA_API_KEY 环境变量")
    
    url = f"{GAMMA_API_BASE_URL}/generations/{generation_id}"
    headers = {
        "X-API-KEY": GAMMA_API_KEY,
        "Content-Type": "application/json"
    }
    
    response = requests.get(url, headers=headers)
    
    if response.status_code != 200:
        print(f"查询状态失败: {response.status_code}")
        print(f"错误信息: {response.text}")
        return None, None
    
    result = response.json()
    status = result.get("status", "unknown")
    return status, result


def wait_for_completion(generation_id, max_wait_time=300, check_interval=5):
    """
    等待生成完成
    
    参数:
    - generation_id: 生成任务 ID
    - max_wait_time: 最大等待时间（秒）
    - check_interval: 检查间隔（秒）
    
    返回: (success, result_data)
    """
    start_time = time.time()
    
    while time.time() - start_time < max_wait_time:
        status, result = check_generation_status(generation_id)
        
        if status is None:
            return False, None
        
        print(f"当前状态: {status}")
        
        if status == "completed":
            print("生成完成！")
            return True, result
        elif status == "failed":
            print("生成失败！")
            return False, result
        elif status in ["pending", "processing"]:
            print(f"等待中... ({int(time.time() - start_time)}秒)")
            time.sleep(check_interval)
        else:
            print(f"未知状态: {status}")
            time.sleep(check_interval)
    
    print(f"超时：等待时间超过 {max_wait_time} 秒")
    return False, None


def download_via_api(gamma_url, generation_id, output_path, export_format="pdf"):
    """
    尝试通过 API 下载导出的文件
    
    参数:
    - gamma_url: Gamma 演示文稿的 URL
    - generation_id: 生成任务 ID
    - output_path: 输出文件路径
    - export_format: 导出格式 ("pdf" 或 "pptx")
    
    返回: (success, method_used)
    """
    headers = {
        "X-API-KEY": GAMMA_API_KEY,
        "Accept": f"application/{export_format}" if export_format == "pdf" else "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    }
    
    import re
    doc_id = None
    
    # 从 URL 中提取文档 ID
    doc_id_match = re.search(r'/docs/([^/?]+)', gamma_url)
    if doc_id_match:
        doc_id = doc_id_match.group(1)
    
    # 尝试多种可能的 API 端点（根据官方文档可能支持的格式）
    if doc_id:
        export_urls = [
            f"{GAMMA_API_BASE_URL}/docs/{doc_id}/export/{export_format}",
            f"{GAMMA_API_BASE_URL}/generations/{generation_id}/export/{export_format}",
            f"https://gamma.app/api/export/{export_format}/{doc_id}",
        ]
    else:
        export_urls = [
            f"{GAMMA_API_BASE_URL}/generations/{generation_id}/export/{export_format}",
        ]
    
    for export_url in export_urls:
        try:
            print(f"尝试 API 端点: {export_url}")
            response = requests.get(export_url, headers=headers, stream=True, timeout=60)
            
            if response.status_code == 200:
                # 检查内容类型或文件头
                content_type = response.headers.get('Content-Type', '')
                
                # 检查是否是目标格式
                is_valid = False
                if export_format == "pdf":
                    is_valid = (
                        'pdf' in content_type.lower() or 
                        response.content[:4] == b'%PDF'
                    )
                elif export_format == "pptx":
                    is_valid = (
                        'presentation' in content_type.lower() or
                        'pptx' in content_type.lower() or
                        response.content[:2] == b'PK'  # PPTX 是 ZIP 格式
                    )
                
                if is_valid:
                    with open(output_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                    print(f"{export_format.upper()} 已通过 API 保存到: {output_path}")
                    return True, "api"
                else:
                    print(f"响应不是 {export_format.upper()} 格式，Content-Type: {content_type}")
                    continue
            elif response.status_code == 202:
                print(f"导出请求已接受，但需要等待处理 (状态码 202)")
                continue
            else:
                print(f"API 请求失败 (状态码 {response.status_code})")
                if response.text:
                    print(f"错误信息: {response.text[:200]}")
                continue
        except Exception as e:
            print(f"API 请求异常: {e}")
            continue
    
    return False, None


def download_via_browser(gamma_url, output_path, export_format="pdf", headless=False):
    """
    使用浏览器自动化下载导出的文件
    
    参数:
    - gamma_url: Gamma 演示文稿的 URL
    - output_path: 输出文件路径
    - export_format: 导出格式 ("pdf" 或 "pptx")
    - headless: 是否使用无头模式（默认 False，便于调试）
    
    返回: (success, method_used)
    """
    if not SELENIUM_AVAILABLE:
        print("Selenium 未安装，无法使用浏览器自动化")
        print("请运行: pip install selenium webdriver-manager")
        return False, None
    
    try:
        print(f"\n{'='*60}")
        print(f"使用浏览器自动化导出 {export_format.upper()}")
        print(f"{'='*60}")
        
        # 配置 Chrome 选项
        chrome_options = Options()
        if headless:
            chrome_options.add_argument('--headless=new')  # 使用新的 headless 模式
        chrome_options.add_argument('--no-sandbox')
        chrome_options.add_argument('--disable-dev-shm-usage')
        chrome_options.add_argument('--disable-gpu')
        chrome_options.add_argument('--window-size=1920,1080')
        
        # 设置下载目录（使用绝对路径）
        download_dir = str(output_path.parent.absolute())
        print(f"下载目录: {download_dir}")
        
        prefs = {
            "download.default_directory": download_dir,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "safebrowsing.enabled": False,  # 禁用安全浏览以加快下载
            "profile.default_content_setting_values.notifications": 2  # 禁用通知
        }
        chrome_options.add_experimental_option("prefs", prefs)
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)
        
        # 创建 WebDriver
        print("正在启动 Chrome 浏览器...")
        try:
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
        except Exception as e:
            print(f"启动浏览器失败: {e}")
            print("请确保已安装 Chrome 浏览器")
            return False, None
        
        try:
            # 访问 Gamma URL
            print(f"\n访问 Gamma URL: {gamma_url}")
            driver.get(gamma_url)
            
            # 等待页面加载
            print("等待页面加载...")
            wait = WebDriverWait(driver, 30)
            time.sleep(8)  # 等待页面完全加载和渲染
            
            # 尝试查找菜单按钮或更多选项按钮（通常 Gamma 的导出在菜单中）
            print("\n查找导出选项...")
            
            # 多种可能的菜单/导出按钮选择器
            menu_selectors = [
                # 通过文本查找
                (By.XPATH, "//button[contains(text(), 'Export')]"),
                (By.XPATH, "//button[contains(text(), '导出')]"),
                (By.XPATH, "//button[contains(text(), 'Download')]"),
                (By.XPATH, "//button[contains(text(), '下载')]"),
                # 通过 aria-label
                (By.CSS_SELECTOR, "button[aria-label*='Export']"),
                (By.CSS_SELECTOR, "button[aria-label*='导出']"),
                (By.CSS_SELECTOR, "button[aria-label*='Menu']"),
                (By.CSS_SELECTOR, "button[aria-label*='菜单']"),
                # 通过 data 属性
                (By.CSS_SELECTOR, "[data-testid*='export']"),
                (By.CSS_SELECTOR, "[data-testid*='menu']"),
                # 通过类名
                (By.CSS_SELECTOR, ".export-button"),
                (By.CSS_SELECTOR, ".menu-button"),
                # 查找包含三个点的菜单按钮
                (By.CSS_SELECTOR, "button[aria-label*='More']"),
                (By.XPATH, "//button[contains(@aria-label, 'More')]"),
            ]
            
            export_button = None
            for by, selector in menu_selectors:
                try:
                    elements = driver.find_elements(by, selector)
                    if elements:
                        # 检查元素是否可见和可点击
                        for elem in elements:
                            if elem.is_displayed() and elem.is_enabled():
                                export_button = elem
                                print(f"找到按钮: {selector}")
                                break
                        if export_button:
                            break
                except:
                    continue
            
            if not export_button:
                # 尝试查找页面上的所有按钮，寻找可能的导出按钮
                print("尝试查找页面上的所有按钮...")
                try:
                    all_buttons = driver.find_elements(By.TAG_NAME, "button")
                    for btn in all_buttons:
                        try:
                            text = btn.text.lower()
                            aria_label = btn.get_attribute("aria-label") or ""
                            if any(keyword in text or keyword in aria_label.lower() 
                                   for keyword in ['export', '导出', 'download', '下载', 'menu', '菜单']):
                                if btn.is_displayed():
                                    export_button = btn
                                    print(f"找到可能的导出按钮: {btn.text or aria_label}")
                                    break
                        except:
                            continue
                except:
                    pass
            
            if export_button:
                try:
                    # 滚动到按钮位置
                    driver.execute_script("arguments[0].scrollIntoView(true);", export_button)
                    time.sleep(1)
                    
                    # 点击按钮
                    print("点击导出/菜单按钮...")
                    export_button.click()
                    time.sleep(3)  # 等待菜单打开
                except Exception as e:
                    print(f"点击按钮失败: {e}")
                    # 尝试使用 JavaScript 点击
                    try:
                        driver.execute_script("arguments[0].click();", export_button)
                        time.sleep(3)
                    except:
                        pass
            else:
                print("未找到导出按钮，尝试使用键盘快捷键...")
                from selenium.webdriver.common.keys import Keys
                from selenium.webdriver.common.action_chains import ActionChains
                
                # 尝试 Ctrl+E 或其他快捷键
                body = driver.find_element(By.TAG_NAME, "body")
                actions = ActionChains(driver)
                actions.send_keys(Keys.ESCAPE).perform()  # 先按 ESC 关闭可能的弹窗
                time.sleep(1)
                actions.key_down(Keys.CONTROL).send_keys('e').key_up(Keys.CONTROL).perform()
                time.sleep(3)
            
            # 查找导出格式选项（PDF 或 PPTX）
            print(f"\n查找 {export_format.upper()} 导出选项...")
            format_selectors = [
                (By.XPATH, f"//button[contains(text(), '{export_format.upper()}')]"),
                (By.XPATH, f"//a[contains(text(), '{export_format.upper()}')]"),
                (By.XPATH, f"//div[contains(text(), '{export_format.upper()}')]"),
                (By.XPATH, f"//*[contains(text(), 'PDF')]" if export_format == "pdf" else f"//*[contains(text(), 'PPTX')]"),
                (By.XPATH, f"//*[contains(text(), 'PowerPoint')]" if export_format == "pptx" else None),
            ]
            
            format_button = None
            for by, selector in format_selectors:
                if selector is None:
                    continue
                try:
                    elements = driver.find_elements(by, selector)
                    for elem in elements:
                        if elem.is_displayed():
                            format_button = elem
                            print(f"找到格式选项: {selector}")
                            break
                    if format_button:
                        break
                except:
                    continue
            
            if format_button:
                try:
                    driver.execute_script("arguments[0].scrollIntoView(true);", format_button)
                    time.sleep(1)
                    print(f"点击 {export_format.upper()} 选项...")
                    format_button.click()
                    time.sleep(2)
                except Exception as e:
                    print(f"点击格式选项失败: {e}")
                    try:
                        driver.execute_script("arguments[0].click();", format_button)
                        time.sleep(2)
                    except:
                        pass
            
            # 等待下载完成
            print(f"\n等待 {export_format.upper()} 下载完成...")
            print("提示: 如果下载未自动开始，请在浏览器中手动点击导出按钮")
            
            max_wait = 120  # 增加等待时间到 2 分钟
            waited = 0
            start_time = time.time()
            
            while waited < max_wait:
                # 检查目标文件是否存在
                if output_path.exists():
                    file_size = output_path.stat().st_size
                    if file_size > 0:
                        print(f"\n✓ {export_format.upper()} 已成功下载到: {output_path}")
                        print(f"  文件大小: {file_size / 1024:.2f} KB")
                        return True, "browser"
                
                # 检查是否有临时下载文件（Chrome 的 .crdownload 文件）
                temp_files = list(output_path.parent.glob(f"*.crdownload"))
                if temp_files:
                    print(f"  下载中... ({waited}秒)")
                
                # 查找最新的对应格式文件
                all_files = list(output_path.parent.glob(f"*.{export_format}"))
                if all_files:
                    latest_file = max(all_files, key=lambda p: p.stat().st_mtime)
                    file_age = time.time() - latest_file.stat().st_mtime
                    if file_age < 30:  # 最近30秒内创建或修改
                        if latest_file != output_path:
                            print(f"  发现新下载的文件: {latest_file.name}")
                            latest_file.rename(output_path)
                            print(f"\n✓ {export_format.upper()} 已下载到: {output_path}")
                            return True, "browser"
                
                time.sleep(2)
                waited += 2
                
                # 每 10 秒显示一次进度
                if waited % 10 == 0:
                    print(f"  等待中... ({waited}/{max_wait}秒)")
            
            print(f"\n✗ 下载超时（等待 {max_wait} 秒）")
            print("提示:")
            print("1. 检查浏览器窗口，可能需要手动点击导出")
            print("2. 检查下载目录是否有文件")
            print("3. 尝试不使用 headless 模式（修改代码中的 headless=False）")
            
            # 如果不是 headless 模式，保持浏览器打开以便用户手动操作
            if not headless:
                print("\n浏览器将保持打开 30 秒，您可以手动完成导出...")
                time.sleep(30)
            
            return False, None
            
        finally:
            if headless or not output_path.exists():
                driver.quit()
            else:
                print("\n浏览器将保持打开，您可以继续操作...")
                print("完成后请手动关闭浏览器")
            
    except Exception as e:
        print(f"\n✗ 浏览器自动化失败: {e}")
        import traceback
        traceback.print_exc()
        return False, None


def download_pdf(gamma_url, generation_id, output_path, use_browser=False):
    """
    从 Gamma URL 下载 PDF 文件
    
    根据官方文档，尝试多种方法导出 PDF：
    1. 优先尝试 API 端点
    2. 如果 API 失败，使用浏览器自动化（如果启用）
    
    参数:
    - gamma_url: Gamma 演示文稿的 URL
    - generation_id: 生成任务 ID（用于 API 调用）
    - output_path: 输出文件路径
    - use_browser: 是否在 API 失败后使用浏览器自动化
    """
    print("\n开始导出 PDF...")
    
    # 方法1: 尝试通过 API 下载
    print("方法1: 尝试通过 API 导出...")
    success, method = download_via_api(gamma_url, generation_id, output_path, "pdf")
    
    if success:
        return True
    
    # 方法2: 如果 API 失败且允许使用浏览器，尝试浏览器自动化
    if use_browser or SELENIUM_AVAILABLE:
        if SELENIUM_AVAILABLE:
            print("\n方法2: API 失败，尝试使用浏览器自动化导出...")
            # 默认不使用 headless 模式，便于调试和手动操作
            headless_mode = os.getenv("BROWSER_HEADLESS", "false").lower() == "true"
            success, method = download_via_browser(gamma_url, output_path, "pdf", headless=headless_mode)
            if success:
                return True
        else:
            print("\nSelenium 未安装，无法使用浏览器自动化")
            print("请运行: pip install selenium webdriver-manager")
    
    # 所有方法都失败
    print("\n所有导出方法均失败")
    print(f"Gamma URL: {gamma_url}")
    print(f"Generation ID: {generation_id}")
    print("\n提示:")
    print("1. 在浏览器中访问上述 Gamma URL")
    print("2. 在网页界面中手动导出 PDF")
    print("3. 安装 Selenium 以启用浏览器自动化: pip install selenium webdriver-manager")
    print("4. 使用 --browser 参数强制使用浏览器自动化")
    return False


def download_pptx(gamma_url, generation_id, output_path, use_browser=False):
    """
    从 Gamma URL 下载 PPTX 文件
    
    参数:
    - gamma_url: Gamma 演示文稿的 URL
    - generation_id: 生成任务 ID
    - output_path: 输出文件路径
    - use_browser: 是否使用浏览器自动化
    """
    print("\n开始导出 PPTX...")
    
    # 方法1: 尝试通过 API 下载
    print("方法1: 尝试通过 API 导出...")
    success, method = download_via_api(gamma_url, generation_id, output_path, "pptx")
    
    if success:
        return True
    
    # 方法2: 使用浏览器自动化
    if use_browser or (not use_browser and SELENIUM_AVAILABLE):
        if SELENIUM_AVAILABLE:
            print("\n方法2: API 失败，尝试使用浏览器自动化导出...")
            success, method = download_via_browser(gamma_url, output_path, "pptx", headless=True)
            if success:
                return True
        else:
            print("\nSelenium 未安装，无法使用浏览器自动化")
    
    print("\n所有导出方法均失败")
    print("提示: PPTX 导出可能需要通过网页界面手动完成")
    return False


def process_file(input_file_path, force_regenerate=False):
    """
    处理单个文件：提取文本、生成 PPT、下载 PDF
    
    参数:
    - input_file_path: 输入文件路径
    - force_regenerate: 是否强制重新生成（即使已存在记录）
    """
    input_path = Path(input_file_path)
    
    if not input_path.exists():
        print(f"文件不存在: {input_path}")
        return False
    
    print(f"\n处理文件: {input_path.name}")
    
    # 检查是否已经生成过
    if not force_regenerate:
        exists, record = check_existing_generation(input_path)
        if exists:
            print(f"\n发现已存在的生成记录:")
            print(f"  生成 ID: {record.get('generation_id', 'N/A')}")
            print(f"  Gamma URL: {record.get('gamma_url', 'N/A')}")
            print(f"  状态: {record.get('status', 'N/A')}")
            print(f"  创建时间: {record.get('created_at', 'N/A')}")
            
            gamma_url = record.get('gamma_url')
            if gamma_url:
                print(f"\n使用已存在的生成结果:")
                print(f"  URL: {gamma_url}")
                
                # 尝试下载 PDF（如果之前没有下载成功）
                output_filename = input_path.stem + "_gamma_presentation.pdf"
                output_path = OUTPUT_DIR / output_filename
                
                if not output_path.exists():
                    print("尝试下载 PDF...")
                    generation_id = record.get('generation_id')
                    download_pdf(gamma_url, generation_id, output_path)
                else:
                    print(f"PDF 文件已存在: {output_path}")
                
                return True
            else:
                print("警告: 记录中没有 Gamma URL，将重新生成")
    
    # 提取文本内容和图片链接
    image_urls = []
    if input_path.suffix.lower() == ".docx":
        result = extract_text_from_docx(input_path)
        if result is None:
            print("未能提取文本内容")
            return False
        text_content, image_urls = result
    elif input_path.suffix.lower() == ".pdf":
        text_content = extract_text_from_pdf(input_path)
        # PDF 中的图片链接提取（简单版本）
        if text_content:
            import re
            urls = re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', text_content)
            image_urls = [url for url in urls if any(url.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg'])]
    else:
        print(f"不支持的文件格式: {input_path.suffix}")
        return False
    
    if not text_content or not text_content.strip():
        print("未能提取到文本内容")
        return False
    
    print(f"已提取文本内容，长度: {len(text_content)} 字符")
    
    # 处理图片 URL（根据官方文档，可以在 inputText 中直接插入图片 URL）
    if image_urls:
        print(f"找到 {len(image_urls)} 个图片链接")
        # 在文本末尾添加图片 URL（根据官方文档，Gamma 会自动识别并处理）
        image_section = "\n\n---\n\n# 图片资源\n"
        for i, img_url in enumerate(image_urls[:20], 1):  # 限制最多20张图片
            image_section += f"{img_url}\n"
        text_content += image_section
        print("已将图片链接添加到输入文本中")
    
    # 检查文本长度（Gamma API 限制约 100,000 tokens）
    if len(text_content) > 400000:  # 粗略估算，1 token ≈ 4 字符
        print("警告: 文本内容可能超过 API 限制，将截断")
        text_content = text_content[:400000]
    
    # 调用 Gamma API 生成演示文稿
    # 如果设置了 GAMMA_THEME_ID 环境变量则使用，否则使用工作区默认主题
    # 根据官方文档，在 additionalInstructions 中添加关于图片的提示
    additional_instructions = "Make the presentation professional and engaging"
    if image_urls:
        additional_instructions += ". Please include and use the image URLs provided in the input text. Display images appropriately in the presentation slides."
    
    generation_id = generate_presentation(
        input_text=text_content,
        theme_id=GAMMA_THEME_ID,  # 如果为 None，API 将使用工作区默认主题
        additional_instructions=additional_instructions,
        has_image_urls=len(image_urls) > 0  # 传递是否有图片 URL 的标志
    )
    
    if not generation_id:
        return False
    
    # 等待生成完成
    success, result = wait_for_completion(generation_id)
    
    if not success:
        return False
    
    # 获取 Gamma URL
    gamma_url = result.get("gammaUrl") or result.get("url")
    
    if not gamma_url:
        print("未获取到 Gamma URL")
        return False
    
    print(f"Gamma URL: {gamma_url}")
    
    # 保存生成记录
    add_generation_record(input_path, generation_id, gamma_url, "completed")
    print("已保存生成记录")
    
    # 生成输出文件名
    output_filename = input_path.stem + "_gamma_presentation.pdf"
    output_path = OUTPUT_DIR / output_filename
    
    # 下载 PDF（传入 generation_id 用于 API 调用）
    # 检查是否使用浏览器自动化
    use_browser = os.getenv("USE_BROWSER_EXPORT", "false").lower() == "true"
    download_success = download_pdf(gamma_url, generation_id, output_path, use_browser=use_browser)
    
    # 更新记录（记录下载状态）
    if download_success:
        update_generation_record(generation_id, pdf_downloaded=True, pdf_path=str(output_path))
    else:
        update_generation_record(generation_id, pdf_downloaded=False)
    
    return download_success


def main():
    """主函数"""
    import sys
    
    print("=" * 50)
    print("Gamma API 演示文稿生成工具")
    print("=" * 50)
    
    # 检查命令行参数
    force_regenerate = False
    use_browser = False
    
    if len(sys.argv) > 1:
        for arg in sys.argv[1:]:
            if arg == "--list" or arg == "-l":
                list_generations()
                return
            elif arg == "--force" or arg == "-f":
                force_regenerate = True
            elif arg == "--browser" or arg == "-b":
                use_browser = True
            else:
                print(f"未知参数: {arg}")
                print("\n用法:")
                print("  python main.py              # 正常处理（跳过已生成的）")
                print("  python main.py --force     # 强制重新生成")
                print("  python main.py --list      # 列出所有生成记录")
                print("  python main.py --browser   # 使用浏览器自动化导出")
                print("  python main.py --force --browser  # 强制重新生成并使用浏览器导出")
                return
    
    # 设置浏览器导出选项
    if use_browser:
        os.environ["USE_BROWSER_EXPORT"] = "true"
    
    # 检查 API Key
    if not GAMMA_API_KEY:
        print("错误: 请设置 GAMMA_API_KEY 环境变量")
        print("可以在 .env 文件中设置，或使用环境变量")
        return
    
    # 查找 dataset 目录下的文件
    docx_files = list(DATASET_DIR.glob("*.docx"))
    pdf_files = list(DATASET_DIR.glob("*.pdf"))
    
    # 优先处理 .docx 文件（原始邮件数据）
    if docx_files:
        print(f"\n找到 {len(docx_files)} 个 .docx 文件")
        for docx_file in docx_files:
            if "原始邮件数据" in docx_file.name:
                process_file(docx_file, force_regenerate=force_regenerate)
                break
    elif pdf_files:
        print(f"\n找到 {len(pdf_files)} 个 PDF 文件")
        # 如果 Gamma 不支持 PDF，这里可以提取文本后处理
        print("注意: Gamma API 不支持直接 PDF 输入，将提取文本内容")
        for pdf_file in pdf_files:
            process_file(pdf_file, force_regenerate=force_regenerate)
            break
    else:
        print("未找到可处理的文件（.docx 或 .pdf）")
    
    print("\n处理完成！")
    print(f"\n提示: 使用 'python main.py --list' 查看所有生成记录")


if __name__ == "__main__":
    main()

